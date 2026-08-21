"""Flask application entry point for Workbook Automation Tool."""

from __future__ import annotations

from pathlib import Path
import re
import shutil
import time
import webbrowser
from uuid import uuid4

from flask import Flask, jsonify, render_template, request, send_file
from werkzeug.utils import secure_filename

from config import Settings, TEAM_CONFIG
from db import DatabaseConfigurationError, OracleDatabase, TaskNotFoundError
from services.pii_service import PIIScanner
from services.workbook_service import WorkbookError, WorkbookProcessor
from services.dsr_service import DSRConfigurationError, DSRService
from services.package_service import package_directory, retain_workbook as retain_package_workbook, safe_token


def create_app(settings: Settings | None = None) -> Flask:
    settings = settings or Settings.from_environment()
    app = Flask(__name__)
    database = OracleDatabase(settings)
    processor = WorkbookProcessor()
    pii_scanner = PIIScanner()
    dsr_service = DSRService(database, settings.dsr_control_workbook)
    workbook_sessions: dict[str, dict] = {}
    monitor_sessions: dict[str, dict] = {}

    instructions_root = Path(__file__).resolve().parent / "sheet_instructions"

    def instruction_path(team: str) -> Path | None:
        folder = {"COPS": "cops", "MQA": "mqa", "Channel Marketing": "channel_marketing"}.get(team)
        return instructions_root / folder / "sheet_instructions.txt" if folder else None

    def task_value(task: object, key: str, default=None):
        """Read task metadata from either a TaskRecord or a legacy dict."""
        if isinstance(task, dict):
            return task.get(key, default)
        return getattr(task, key, default)

    def retain_workbook(source: Path, original_name: str | None = None) -> tuple[str, Path]:
        session_id = uuid4().hex
        session_dir = settings.download_dir / "sessions" / session_id
        session_dir.mkdir(parents=True, exist_ok=True)
        retained = session_dir / source.name
        shutil.copy2(source, retained)
        workbook_sessions[session_id] = {
            "source": retained,
            "temporary": None,
            "original_name": original_name or source.name,
        }
        return session_id, retained

    def build_session_payload(
        source: Path, workfront_content: dict, upload_source: str, original_filename: str | None = None
    ) -> dict:
        processing_source = source
        converted = None
        if source.suffix.lower() == ".xls":
            import pandas as pd
            settings.download_dir.mkdir(parents=True, exist_ok=True)
            converted = settings.download_dir / f"converted-{uuid4().hex}-{source.stem}.xlsx"
            try:
                sheets = pd.read_excel(source, sheet_name=None, engine="xlrd")
            except ImportError as error:
                raise WorkbookError("Legacy .xls support is not installed on the server.") from error
            with pd.ExcelWriter(converted, engine="openpyxl") as writer:
                for sheet_name, frame in sheets.items():
                    frame.to_excel(writer, sheet_name=str(sheet_name)[:31], index=False)
            processing_source = converted
        try:
            display_name = original_filename or source.name
            session_id, retained = retain_workbook(processing_source, display_name)
            result = processor.build_payload(retained, workfront_content)
            result.update(
                status="ready",
                upload_source=upload_source,
                workbook_session_id=session_id,
                original_filename=display_name,
            )
            return result
        finally:
            if converted:
                converted.unlink(missing_ok=True)

    def mask_sensitive_text(value: str) -> tuple[str, dict[str, int]]:
        from collections import Counter
        counts: Counter[str] = Counter()
        masked = value
        for label in pii_scanner.SCAN_ORDER:
            for pattern in pii_scanner.PII_PATTERNS[label]:
                def replace(match, sensitive_type=label):
                    counts[sensitive_type] += 1
                    return "[MASKED]"
                masked = pattern.sub(replace, masked)
        return masked, dict(counts)

    @app.get("/")
    def index():
        return render_template("index.html", cops_gpt_url=settings.cops_gpt_url)

    @app.get("/health")
    def health():
        return {"status": "ok"}

    @app.get("/team_config")
    def team_config():
        return jsonify(TEAM_CONFIG)

    @app.get("/sheet_instructions/<team>")
    def sheet_instructions(team: str):
        path = instruction_path(team)
        if not path or not path.is_file():
            return jsonify(error="Sheet instructions were not found for this team."), 404
        return jsonify(team=team, instructions=path.read_text(encoding="utf-8"))

    @app.post("/cops/prepare")
    def prepare_cops():
        payload = request.get_json(silent=True) or {}
        project = str(payload.get("project_reference", "")).strip()
        task_number = str(payload.get("task_number", "")).strip()
        if not project or not task_number:
            return jsonify(error="Project reference and task number are required."), 400
        if len(project) > 100 or len(task_number) > 100:
            return jsonify(error="Project reference and task number must be 100 characters or fewer."), 400
        target = package_directory(settings.monitored_download_dir, project, task_number)
        try:
            result = dsr_service.generate_dsr(
                project, target / "DSR_Report.docx", task_number=task_number
            )
            task = database.get_task(project, task_number)
            return jsonify(
                status="dsr_ready",
                dsr_status="ready",
                workbook_status="available" if task_value(task, "workbook_url") else "manual_required",
                task_url=task_value(task, "task_url"),
                workbook_url=task_value(task, "workbook_url"),
                expected_filename=task_value(task, "filename", ""),
                dsr=result.public_data,
                message="DSR generated. Preparing the workbook.",
            )
        except TaskNotFoundError as error:
            return jsonify(error=str(error)), 404
        except (DatabaseConfigurationError, DSRConfigurationError):
            app.logger.warning("COPS backend configuration is incomplete")
            return jsonify(error="DSR generation failed. Please retry or contact the administrator."), 422
        except Exception:
            app.logger.exception("COPS preparation failed")
            return jsonify(error="Unable to prepare the COPS validation package."), 500

    @app.get("/cops/dsr")
    def download_cops_dsr():
        project = str(request.args.get("project_reference", "")).strip()
        task_number = str(request.args.get("task_number", "")).strip()
        path = package_directory(settings.monitored_download_dir, project, task_number) / "DSR_Report.docx"
        if not path.is_file():
            return jsonify(error="DSR report was not found."), 404
        return send_file(
            path,
            as_attachment=True,
            download_name=f"DSR_PRN-{safe_token(project)}_Task-{safe_token(task_number)}.docx",
        )

    @app.get("/cops/dsr/preview")
    def preview_cops_dsr():
        project = str(request.args.get("project_reference", "")).strip()
        task_number = str(request.args.get("task_number", "")).strip()
        path = package_directory(settings.monitored_download_dir, project, task_number) / "DSR_Report.docx"
        if not path.is_file():
            return jsonify(error="DSR report was not found."), 404
        from docx import Document
        document = Document(path)
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        tables = [[[cell.text for cell in row.cells] for row in table.rows] for table in document.tables]
        return jsonify(filename=path.name, paragraphs=paragraphs, tables=tables)

    @app.post("/cops/data_safety")
    def cops_data_safety():
        payload = request.get_json(silent=True) or {}
        session_id = str(payload.get("workbook_session_id", "")).strip()
        session = workbook_sessions.get(session_id)
        project = str(payload.get("project_reference", "")).strip()
        task_number = str(payload.get("task_number", "")).strip()
        dsr_path = package_directory(settings.monitored_download_dir, project, task_number) / "DSR_Report.docx"
        instructions = instruction_path("COPS")
        try:
            if not session_id or not session:
                return jsonify(error="The workbook session expired. Fetch or upload the workbook again."), 409
            workbook_path = Path(session.get("source", ""))
            if not workbook_path.is_file():
                return jsonify(error="The retained workbook is unavailable. Fetch or upload it again."), 409
            if not dsr_path.is_file():
                return jsonify(error="The DSR report is unavailable. Generate it again before retrying."), 409
            if not instructions or not instructions.is_file():
                return jsonify(error="COPS sheet instructions are unavailable."), 409

            from collections import Counter
            from docx import Document
            from openpyxl import load_workbook
            original_name = session.get("original_name") or workbook_path.name
            original_extension = Path(original_name).suffix.lower() or workbook_path.suffix.lower()
            masked_workbook_name = f"WB_Masked_{Path(original_name).stem}{original_extension}"
            masked_dsr_name = f"DSR_Masked_PRN-{safe_token(project)}_Task-{safe_token(task_number)}{dsr_path.suffix.lower()}"
            masked_workbook_path = workbook_path.parent / masked_workbook_name
            masked_dsr_path = workbook_path.parent / masked_dsr_name

            workbook = load_workbook(workbook_path, keep_vba=workbook_path.suffix.lower() == ".xlsm")
            details = []
            total_counts: Counter[str] = Counter()
            affected_sheets = set()
            masked_preview = []
            try:
                for sheet in workbook.worksheets:
                    sheet_counts: Counter[str] = Counter()
                    for row in sheet.iter_rows():
                        for cell in row:
                            if not isinstance(cell.value, str):
                                continue
                            masked, counts = mask_sensitive_text(cell.value)
                            if counts:
                                cell.value = masked
                                sheet_counts.update(counts)
                    if sheet_counts:
                        affected_sheets.add(sheet.title)
                        total_counts.update(sheet_counts)
                        for sensitive_type, count in sorted(sheet_counts.items()):
                            details.append({"sheet": sheet.title, "type": sensitive_type, "count": count, "status": "Masked"})
                        preview_rows = [
                            [None if cell.value is None else str(cell.value) for cell in row[:20]]
                            for row in sheet.iter_rows(min_row=1, max_row=min(sheet.max_row, 50))
                        ]
                        masked_preview.append({
                            "sheet": sheet.title,
                            "categories": sorted(sheet_counts),
                            "rows": preview_rows,
                        })
                workbook.save(masked_workbook_path)
            finally:
                workbook.close()

            document = Document(dsr_path)
            dsr_counts: Counter[str] = Counter()
            for paragraph in document.paragraphs:
                masked, counts = mask_sensitive_text(paragraph.text)
                if counts:
                    paragraph.text = masked
                    dsr_counts.update(counts)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        masked, counts = mask_sensitive_text(cell.text)
                        if counts:
                            cell.text = masked
                            dsr_counts.update(counts)
            document.save(masked_dsr_path)
            total_counts.update(dsr_counts)
            if dsr_counts:
                for sensitive_type, count in sorted(dsr_counts.items()):
                    details.append({"sheet": "DSR Report", "type": sensitive_type, "count": count, "status": "Masked"})
            session["masked_workbook"] = masked_workbook_path
            session["masked_dsr"] = masked_dsr_path
            session["masked_workbook_name"] = masked_workbook_name
            session["masked_dsr_name"] = masked_dsr_name
            return jsonify(
                status="masked",
                has_pii=bool(total_counts),
                detected_count=sum(total_counts.values()),
                masked_count=sum(total_counts.values()),
                affected_sheets=len(affected_sheets),
                types=dict(sorted(total_counts.items())),
                details=details,
                large_masked_data=sum(total_counts.values()) >= 10 and bool(masked_preview),
                masked_preview=masked_preview,
                workbook_filename=masked_workbook_name,
                dsr_filename=masked_dsr_name,
            )
        except (OSError, ValueError, KeyError) as error:
            app.logger.warning("COPS Data Safety input failure: %s", error)
            return jsonify(error="COPS Data Safety could not read the workbook or DSR. Retry after regenerating them."), 422
        except Exception:
            app.logger.exception("COPS Data Safety failed")
            return jsonify(error="COPS Data Safety could not be completed. Please retry."), 500

    @app.get("/cops/masked/<session_id>/<file_kind>")
    def download_cops_masked_file(session_id: str, file_kind: str):
        session = workbook_sessions.get(session_id)
        key = {"workbook": "masked_workbook", "dsr": "masked_dsr"}.get(file_kind)
        name_key = {"workbook": "masked_workbook_name", "dsr": "masked_dsr_name"}.get(file_kind)
        path = Path(session.get(key, "")) if session and key else None
        if not path or not path.is_file():
            return jsonify(error="Masked file is not ready. Complete Data Safety first."), 409
        return send_file(path, as_attachment=True, download_name=session.get(name_key) or path.name)

    @app.post("/scan_pii")
    def scan_pii():
        payload = request.get_json(silent=True) or {}
        prompt = payload.get("prompt", "")
        sections = payload.get("sections")
        worksheet = str(payload.get("worksheet", "")).strip()
        if not isinstance(prompt, str):
            return jsonify(error="prompt must be a string."), 400
        if sections is not None and (
            not isinstance(sections, dict)
            or not all(isinstance(key, str) and isinstance(value, str) for key, value in sections.items())
        ):
            return jsonify(error="sections must be an object containing text values."), 400
        total_length = len(prompt) + sum(len(value) for value in (sections or {}).values())
        if total_length > 1_000_000:
            return jsonify(error="prompt must be 1,000,000 characters or fewer."), 413
        if sections:
            return jsonify(pii_scanner.scan_sections(sections, worksheet))
        return jsonify(pii_scanner.scan(prompt))

    @app.post("/fetch_workbook")
    def fetch_workbook():
        payload = request.get_json(silent=True) or {}
        project_reference = str(payload.get("project_reference", "")).strip()
        task_number = str(payload.get("task_number", "")).strip()
        if not project_reference or not task_number:
            return jsonify(error="project_reference and task_number are required."), 400
        if len(project_reference) > 100 or len(task_number) > 100:
            return jsonify(error="Project reference and task number must be 100 characters or fewer."), 400

        try:
            task = database.get_task(project_reference, task_number)
            workbook_url = task_value(task, "workbook_url")
            filename = task_value(task, "filename", "")
            task_url = task_value(task, "task_url")
            if not workbook_url or not filename:
                return jsonify(
                    status="manual_required",
                    message=(
                        "Workbook download link not found. Please download the "
                        "workbook manually using the task link below."
                    ),
                    task_url=task_url,
                    expected_filename=filename,
                )
            if not webbrowser.open(workbook_url):
                return jsonify(
                    status="manual_required",
                    message=(
                        "The workbook download could not be opened automatically. "
                        "Please use the Workfront task link to download it."
                    ),
                    task_url=task_url,
                    expected_filename=filename,
                )
            return jsonify(
                status="monitor_required",
                message=(
                    "The workbook download was opened in your browser. "
                    "Monitoring Downloads folder..."
                ),
                task_url=task_url,
                expected_filename=filename,
            )
        except TaskNotFoundError as error:
            return jsonify(error=str(error)), 404
        except (DatabaseConfigurationError, WorkbookError) as error:
            app.logger.warning("Workbook fetch failed: %s", error)
            return jsonify(error=str(error)), 422
        except Exception:
            app.logger.exception("Unexpected workbook fetch failure")
            return jsonify(error="Unable to fetch the workbook. Please contact support if the problem persists."), 500

    @app.post("/lookup_workbook")
    def lookup_workbook():
        payload = request.get_json(silent=True) or {}
        project_reference = str(payload.get("project_reference", "")).strip()
        task_number = str(payload.get("task_number", "")).strip()
        if not project_reference or not task_number:
            return jsonify(error="project_reference and task_number are required."), 400
        try:
            task = database.get_task(project_reference, task_number)
            return jsonify(
                status=(
                    "found"
                    if task_value(task, "workbook_url") and task_value(task, "filename")
                    else "manual_required"
                ),
                task_url=task_value(task, "task_url"),
                filename=task_value(task, "filename", ""),
                message=(
                    None
                    if task_value(task, "workbook_url") and task_value(task, "filename")
                    else (
                        "Workbook download link not found. Please download the "
                        "workbook manually using the task link below."
                    )
                ),
            )
        except TaskNotFoundError as error:
            return jsonify(error=str(error)), 404
        except DatabaseConfigurationError as error:
            return jsonify(error=str(error)), 422

    @app.post("/upload_workbook")
    def upload_workbook():
        uploaded = request.files.get("workbook")
        project_reference = str(request.form.get("project_reference", "")).strip()
        task_number = str(request.form.get("task_number", "")).strip()
        team = str(request.form.get("team", "")).strip()
        if not uploaded or not uploaded.filename:
            return jsonify(error="Select a workbook file to upload."), 400

        filename = secure_filename(uploaded.filename)
        extension = Path(filename).suffix.lower()
        if extension not in {".xls", ".xlsx", ".xlsm"}:
            return jsonify(error="Upload an .xls, .xlsx, or .xlsm workbook."), 400

        settings.download_dir.mkdir(parents=True, exist_ok=True)
        workbook_path = settings.download_dir / f"manual-{uuid4().hex}-{filename}"

        try:
            uploaded.save(workbook_path)
            upload_bytes = workbook_path.stat().st_size
            if upload_bytes > settings.max_workbook_bytes:
                return jsonify(
                    error=f"Workbook exceeds the {settings.max_workbook_bytes // (1024 * 1024)} MB upload limit."
                ), 413
            workfront_content = {}
            if project_reference and (task_number or team == "COPS"):
                try:
                    workfront_content = database.fetch_workfront_content(
                        project_reference, task_number or None
                    )
                except Exception as error:
                    # Workfront metadata enriches the workbook prompt but must not
                    # block a valid manual workbook from entering the shared pipeline.
                    app.logger.warning(
                        "Manual workbook metadata lookup unavailable: team=%s extension=%s bytes=%s error=%s",
                        team or "unknown",
                        extension,
                        upload_bytes,
                        type(error).__name__,
                    )
            result = build_session_payload(workbook_path, workfront_content, "manual", filename)
            result["original_filename"] = filename
            if team == "COPS" and project_reference:
                target = package_directory(settings.monitored_download_dir, project_reference, task_number)
                try:
                    retained = retain_package_workbook(workbook_path, target)
                    result.update(
                        packaged_workbook=retained.name,
                        package_ready=(target / "DSR_Report.docx").is_file(),
                    )
                except OSError as error:
                    app.logger.warning(
                        "COPS package copy unavailable after successful upload: extension=%s bytes=%s error=%s",
                        extension,
                        upload_bytes,
                        type(error).__name__,
                    )
                    result.update(packaged_workbook=None, package_ready=False)
            return jsonify(result)
        except (DatabaseConfigurationError, WorkbookError) as error:
            app.logger.warning(
                "Manual workbook upload failed: team=%s extension=%s error=%s",
                team or "unknown",
                extension,
                type(error).__name__,
            )
            return jsonify(error=str(error)), 422
        except Exception as error:
            app.logger.exception(
                "Unexpected manual workbook upload failure: team=%s extension=%s error=%s",
                team or "unknown",
                extension,
                type(error).__name__,
            )
            return jsonify(error="Unable to store or process the workbook."), 500
        finally:
            workbook_path.unlink(missing_ok=True)

    @app.post("/monitor_workbook/start")
    def start_workbook_monitor():
        payload = request.get_json(silent=True) or {}
        project_reference = str(payload.get("project_reference", "")).strip()
        task_number = str(payload.get("task_number", "")).strip()
        expected_filename = str(payload.get("expected_filename", "")).strip()
        team = str(payload.get("team", "")).strip()
        if not project_reference or (not task_number and team != "COPS"):
            return jsonify(error="Project reference and task number are required."), 400
        monitor_id = uuid4().hex
        monitor_sessions[monitor_id] = {
            "project_reference": project_reference,
            "task_number": task_number,
            "expected_filename": Path(expected_filename).name,
            "started_at": time.time(),
            "cancelled": False,
            "detected_path": None,
            "team": team,
        }
        return jsonify(
            status="monitoring",
            monitor_id=monitor_id,
            timeout_seconds=settings.download_monitor_timeout,
            message="Monitoring Downloads folder...",
        )

    @app.post("/monitor_workbook/poll")
    def poll_workbook_monitor():
        payload = request.get_json(silent=True) or {}
        monitor_id = str(payload.get("monitor_id", ""))
        monitor = monitor_sessions.get(monitor_id)
        if not monitor:
            return jsonify(error="Monitoring session was not found."), 404
        if monitor["cancelled"]:
            monitor_sessions.pop(monitor_id, None)
            return jsonify(status="cancelled", message="Workbook monitoring cancelled.")
        if time.time() - monitor["started_at"] >= settings.download_monitor_timeout:
            monitor_sessions.pop(monitor_id, None)
            return jsonify(status="timeout", message="No matching workbook was detected before timeout.")

        if monitor.get("detected_path"):
            detected = Path(monitor["detected_path"])
            try:
                task = database.get_task(
                    monitor["project_reference"],
                    monitor["task_number"] or None,
                )
                result = build_session_payload(
                    detected, task_value(task, "workfront_content", {}), "detected"
                )
                if monitor.get("team") == "COPS":
                    target = package_directory(
                        settings.monitored_download_dir,
                        monitor["project_reference"], monitor["task_number"] or "auto",
                    )
                    retained = retain_package_workbook(detected, target)
                    result.update(
                        packaged_workbook=retained.name,
                        package_ready=(target / "DSR_Report.docx").is_file(),
                    )
                monitor_sessions.pop(monitor_id, None)
                result["message"] = "Workbook uploaded successfully."
                return jsonify(result)
            except (DatabaseConfigurationError, WorkbookError) as error:
                return jsonify(error=str(error)), 422

        directory = settings.monitored_download_dir
        expected = monitor["expected_filename"].lower()
        expected_stem = Path(expected).stem
        project_token = re.sub(r"[^a-z0-9]", "", monitor["project_reference"].lower())
        task_token = re.sub(r"[^a-z0-9]", "", monitor["task_number"].lower())
        candidates: list[Path] = []
        if directory.is_dir():
            for candidate in directory.iterdir():
                if candidate.suffix.lower() not in {".xls", ".xlsx", ".xlsm"}:
                    continue
                try:
                    if candidate.stat().st_mtime < monitor["started_at"]:
                        continue
                except OSError:
                    continue
                normalized = re.sub(r"[^a-z0-9]", "", candidate.stem.lower())
                filename_matches = bool(
                    expected_stem
                    and (
                        candidate.name.lower() == expected
                        or candidate.stem.lower().startswith(f"{expected_stem} (")
                    )
                )
                identifiers_match = (
                    project_token in normalized
                    and task_token in normalized
                )
                if filename_matches or (not expected and identifiers_match):
                    candidates.append(candidate)

        if not candidates:
            return jsonify(status="monitoring", message="Waiting for workbook download...")

        detected = max(candidates, key=lambda path: path.stat().st_mtime)
        try:
            size_before = detected.stat().st_size
            time.sleep(0.25)
            if detected.stat().st_size != size_before:
                return jsonify(status="monitoring", message="Workbook detected. Waiting for download to finish...")
            monitor["detected_path"] = str(detected)
            return jsonify(
                status="monitoring",
                message="Workbook detected. Uploading workbook...",
            )
        except OSError:
            return jsonify(status="monitoring", message="Waiting for workbook download...")

    @app.post("/monitor_workbook/cancel")
    def cancel_workbook_monitor():
        payload = request.get_json(silent=True) or {}
        monitor = monitor_sessions.get(str(payload.get("monitor_id", "")))
        if monitor:
            monitor["cancelled"] = True
        return jsonify(status="cancelled")

    @app.post("/temporary_workbook")
    def create_temporary_workbook():
        payload = request.get_json(silent=True) or {}
        session_id = str(payload.get("workbook_session_id", ""))
        selected_sheets = payload.get("selected_sheets")
        session = workbook_sessions.get(session_id)
        if not session:
            return jsonify(error="Workbook session was not found."), 404
        if not isinstance(selected_sheets, list) or not selected_sheets:
            return jsonify(error="Select at least one worksheet."), 400
        from openpyxl import load_workbook

        workbook = load_workbook(session["source"], data_only=False)
        try:
            missing = [name for name in selected_sheets if name not in workbook.sheetnames]
            if missing:
                return jsonify(error=f"Unknown worksheet: {missing[0]}"), 400
            for worksheet in list(workbook.worksheets):
                if worksheet.title not in selected_sheets:
                    workbook.remove(worksheet)
            selected_objects = {worksheet.title: worksheet for worksheet in workbook.worksheets}
            workbook._sheets = [selected_objects[name] for name in selected_sheets]
            temporary_name = f"validation-{session_id[:8]}.xlsx"
            temporary_path = session["source"].parent / temporary_name
            workbook.save(temporary_path)
        finally:
            workbook.close()
        prior = session.get("temporary")
        if prior and prior != temporary_path:
            Path(prior).unlink(missing_ok=True)
        session["temporary"] = temporary_path
        return jsonify(
            attachment_id=session_id,
            filename=temporary_name,
            sheets=selected_sheets,
            size=temporary_path.stat().st_size,
            download_url=f"/temporary_workbook/{session_id}",
        )

    @app.get("/temporary_workbook/<session_id>")
    def download_temporary_workbook(session_id: str):
        session = workbook_sessions.get(session_id)
        path = Path(session["temporary"]) if session and session.get("temporary") else None
        if not path or not path.is_file():
            return jsonify(error="Temporary workbook was not found."), 404
        return send_file(path, as_attachment=True, download_name=path.name)

    @app.get("/workbook_session/<session_id>/source")
    def download_session_workbook(session_id: str):
        session = workbook_sessions.get(session_id)
        path = Path(session["source"]) if session and session.get("source") else None
        if not path or not path.is_file():
            return jsonify(error="Workbook session was not found."), 404
        return send_file(path, as_attachment=True, download_name=session.get("original_name") or path.name)

    @app.post("/workbook_session/end")
    def end_workbook_session():
        payload = request.get_json(silent=True) or {}
        session_id = str(payload.get("workbook_session_id", ""))
        session = workbook_sessions.pop(session_id, None)
        if session:
            shutil.rmtree(Path(session["source"]).parent, ignore_errors=True)
        return jsonify(status="ended")

    return app


if __name__ == "__main__":
    create_app().run(host="127.0.0.1", port=5000, debug=True)
