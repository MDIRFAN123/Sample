import os
import re
import sys
import imghdr
from argparse import ArgumentParser
from flask import Blueprint, Flask, request, redirect, url_for, flash, \
    send_from_directory, abort, render_template
from werkzeug.utils import secure_filename
from werkzeug.middleware.proxy_fix import ProxyFix
from flask import Flask, request, render_template_string, send_file, flash
import threading
import time
import oracledb
import docx
import pandas as pd
import numpy as np
import json
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from datetime import datetime
import pytz
import tempfile
import zipfile

app = Flask(__name__)
app.secret_key = "super secret key"
app.wsgi_app = ProxyFix(app.wsgi_app)
app.config.update(PREFERRED_URL_SCHEME='https',
                  project_hosts=[])

bp = Blueprint('main', __name__)

def get_db_connection():
    # Use your actual credentials here
    sso = 'XXXXXXXXXX'
    pw = 'XXXXXXXXXXXXXXXX'
    uri_p3_p4 = 'XXXXXXXXXXXXXX'
    port_p3_p4_p5 = 'XXXXXXXX'
    db_name_p3 = 'XXXXXXXXX'

    oracledb.init_oracle_client()  # Enabling Thick Mode

    connection = oracledb.connect(
        '{sso}/{pw}@{uri}:{port}/{db_name}'.format(
            sso=sso,
            pw=pw,
            uri=uri_p3_p4,
            port=port_p3_p4_p5,
            db_name=db_name_p3
        )
    )

    return connection


def get_or_create_hyperlink_style(d):
    """
    If this document had no hyperlinks so far, the builtin
    Hyperlink style will likely be missing and we need to add it.
    There's no predefined value, different Word versions
    define it differently.
    This version is how Word 2019 defines it in the
    default theme, excluding a theme reference.
    """

    if "Hyperlink" not in d.styles:
        if "Default Character Font" not in d.styles:
            ds = d.styles.add_style(
                "Default Character Font",
                docx.enum.style.WD_STYLE_TYPE.CHARACTER,
            )

            ds.element.set(docx.oxml.shared.qn('w:default'), "1")
            ds.priority = 1
            ds.hidden = True
            ds.unhide_when_used = True

            del ds

        hs.base_style = d.styles["Default Character Font"]
        hs.unhide_when_used = True
        hs.font.color.rgb = docx.shared.RGBColor(0x05, 0x63, 0xC1)
        hs.font.underline = True

    return "Hyperlink"

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url,
        docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True
    )

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

    # Create a new run object (a wrapper over a 'w:r' element)
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'),
        paragraph
    )
    new_run.text = text

    # Set the run's style to the builtin hyperlink style, defining it if necessary
    new_run.style = get_or_create_hyperlink_style(part.document)

    # Alternatively, set the run's formatting explicitly
    # new_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    # new_run.font.underline = True

    # Join all the xml elements together
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)

    return "Hyperlink"


def generate_report(workfront_num):

    # The fixed timezone & timestamp
    eastern = pytz.timezone('US/Eastern')
    current_time = datetime.now(eastern)
    timestamp_string = current_time.strftime("%Y-%m-%d %H:%M:%S")

    # Your helper functions to set background, no wrap
    def set_cell_background(cell, color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def set_no_wrap(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        no_wrap = OxmlElement('w:noWrap')
        tcPr.append(no_wrap)

    # Weight functions, define them here or import/define similarly from your code

    rc_weights = {
        'Data Targeting Criteria': 0.17,
        'Marketing or Servicing': 0.15,
        'Channel': 0.14,
        'Will a start file be provided?': 0.12,
        'Is a control group needed?': 0.11,
        'Campaign Type - Lifecycle': 0.09,
        'Output File Data or MISC fields': 0.08,
        'Number of Segments': 0.06,
        'Non-Mon Required?': 0.04,
        'Marketing Tollgate Approval': 0.03,
        'Creative Agency Name': 0.01
    }

    ps_weights = {
        'Data Targeting Criteria': 0.17,
        'Marketing or Servicing': 0.15,
        'Channel': 0.14,
        'Will a start file be provided?': 0.12,
        'Is a control group needed?': 0.11,
        'Campaign Type - Lifecycle': 0.09,
        'Output File Data or MISC fields': 0.08,
        'Number of Segments': 0.06,
        'Marketing Tollgate Approval': 0.04,
        'Creative Agency Name': 0.03,
        'Non-Mon Required?': 0.01
    }

    zeta_rc_weights = {
        'Data Targeting Criteria': 0.14,
        'Marketing or Servicing': 0.12,
        'ZETA File Name': 0.20,
        'Channel': 0.11,
        'Will a start file be provided?': 0.10,
        'Is a control group needed?': 0.09,
        'Campaign Type - Lifecycle': 0.07,
        'Output File Data or MISC fields': 0.06,
        'Number of Segments': 0.05,
        'Non-Mon Required?': 0.03,
        'Marketing Tollgate Approval': 0.02,
        'Creative Agency Name': 0.01
    }

    zeta_ps_weights = {
        'Data Targeting Criteria': 0.14,
        'Marketing or Servicing': 0.12,
        'ZETA File Name': 0.20,
        'Channel': 0.11,
        'Will a start file be provided?': 0.10,
        'Is a control group needed?': 0.09,
        'Campaign Type - Lifecycle': 0.07,
        'Output File Data or MISC fields': 0.06,
        'Number of Segments': 0.05,
        'Marketing Tollgate Approval': 0.03,
        'Creative Agency Name': 0.02,
        'Non-Mon Required?': 0.01
    }

    def calculate_percentage_rc_col(col):

        if 'zeta' in df.loc['Name', 'Values'].casefold():
            valid_vars = [
                var for var in zeta_rc_weights
                if var in col.index and not pd.isna(col.loc[var])
            ]
            total_weight = sum(
                zeta_rc_weights[var] for var in valid_vars
            )
            return total_weight * 100

        else:
            valid_vars = [
                var for var in rc_weights
                if var in col.index and not pd.isna(col.loc[var])
            ]
            total_weight = sum(
                rc_weights[var] for var in valid_vars
            )
            return total_weight * 100

    def calculate_percentage_ps_col(col):

        if 'zeta' in df.loc['Name', 'Values'].casefold():
            valid_vars = [
                var for var in zeta_ps_weights
                if var in col.index and not pd.isna(col.loc[var])
            ]
            total_weight = sum(
                zeta_ps_weights[var] for var in valid_vars
            )
            return total_weight * 100

        else:
            valid_vars = [
                var for var in ps_weights
                if var in col.index and not pd.isna(col.loc[var])
            ]
            total_weight = sum(
                ps_weights[var] for var in valid_vars
            )
            return total_weight * 100



sql = "SELECT BODY from cmktsch.WF_PROJ WHERE REFERENCENUMBER = :refnum"

with get_db_connection() as conn:
    with conn.cursor() as cursor:
        cursor.execute(sql, {"refnum": workfront_num})
        row = cursor.fetchone()

if not row or row[0] is None:
    raise ValueError(f"No data found for Workfront #: {workfront_num}")

clob = row[0]
clob_text = clob if isinstance(clob, str) else clob.read()

data_dict = json.loads(clob_text)

selected_vars = [
    "Name",
    "IsStatusComplete",
    "DE:Zeta COPS File Name",
    "DE:Channel",
    "DE:Marketing or Servicing",
    "DE:Campaign Type - Lifecycle",
    "DE:Creative Agency Name",
    "DE:Marketing Tollgate Approval",
    "DE:Will a start file be provided?",
    "DE:Data Targeting Criteria",
    "DE:Output File Data or MISC fields",
    "DE:Is a control group needed?",
    "DE:Number of Segments",
    "DE:COPS Campaign Code",
    "DE:Client or Industry calculated",
    "DE:Non-Mon Required?",
    "DE:Open Segmentation Matrix from Project"
]

filtered_dict1 = {
    k: v for k, v in data_dict.items()
    if k in selected_vars
}

filtered_dict2 = {
    k: v for k, v in data_dict["parameterValues"].items()
    if k in selected_vars
}

combined = {**filtered_dict1, **filtered_dict2}

final_dict = {}

for var in selected_vars:
    if var in combined:
        final_dict[var] = combined[var]
    else:
        final_dict[var] = np.nan

df = pd.DataFrame.from_dict(
    final_dict,
    orient="index",
    columns=["Values"]
)

df.index = [
    "Name",
    "Is Status Complete?",
    "ZETA File Name",
    "Channel",
    "Marketing or Servicing",
    "Campaign Type - Lifecycle",
    "Creative Agency Name",
    "Marketing Tollgate Approval",
    "Will a start file be provided?",
    "Data Targeting Criteria",
    "Output File Data or MISC fields",
    "Is a control group needed?",
    "Number of Segments",
    "COPS Campaign Code",
    "Portfolio: Name",
    "Non-Mon Required?",
    "Segmentation Matrix_Link"
]

df.index.name = "Parameters"


# Classify Client

CL = [
    "Vemo Visa",
    "PayPal Co-Brand/PayPal PLCC",
    "Lowe's Commercial",
    "Sam's Club",
    "PayPal Credit",
    "TJX",
    "Cathay Pacific",
    "Verizon",
    "Crate & Barrel All Brands",
    "Fleet Farm (SR6)",
    "ShopH",
    "PayPal All Products",
    "Lifestyle Markets - Multiclients",
    "Acquisitions - B2C",
    "Lowe's Consumer",
    "Dual Card - B2C",
    "Dick's Sporting Goods",
    "At Home",
    "Google",
    "JC Penney",
    "American Signature Furniture",
    "RC Multi-Client",
    "American Eagle",
    "Belk",
    "Harbor Freight Tools",
    "Amazon Consumer",
    "Walgreens",
    "Brand - B2C",
    "Fareportal",
    "PayPal Credit Card",
    "Ebay",
    "J.Crew",
    "Virgin",
    "HSN",
    "Walmart",
    "QVC"
]

if df.loc["Portfolio: Name", "Values"] in CL:
    df.loc["Client"] = {"Values": "RC"}
else:
    df.loc["Client"] = {"Values": "PS"}


availability_percentage_rc = df.apply(
    calculate_percentage_rc_col,
    axis=0
)

availability_percentage_ps = df.apply(
    calculate_percentage_ps_col,
    axis=0
)

if df.loc["Client", "Values"] == "RC":
    df.loc["Availability Percentage"] = int(
        availability_percentage_rc
    )
else:
    df.loc["Availability Percentage"] = int(
        availability_percentage_ps
    )


if "zeta" in df.loc["Name", "Values"].casefold():
    # case-insensitive, better than lower()
    df2 = df.head(df.shape[0] - 3)
else:
    df = df.drop(index="ZETA File Name")
    df2 = df.head(df.shape[0] - 3)


mask = df2["Values"].isnull()

missing_values = (
    df2[mask]
    .fillna("Missing")
    .reset_index()
)

non_missing_values = (
    df2[~mask]
    .fillna("Missing")
    .reset_index()
)


# Fetching the segmentation details

sql_new = """
SELECT
    b.Client_Brand,
    b.seg_group AS Group_Name,
    b.Segment_Name,
    b.Segment_Offer_Description,
    b.Target_Criteria AS Targeting_Criteria,
    b.PRODUCT_TYPE AS Segment_Product_Type,
    b.CHANNEL,
    b.SEG_LANGUAGE AS Language,
    b.control AS Test_Control,
    b.percent_segment AS Segment_Percent,
    b.commcode AS Comm_Code,
    b.Creative_Code,
    b.CIS_MEMO
FROM cmktsch.CC_SEG_MATRIX_HEADER a
RIGHT JOIN cmktsch.CC_SEG_MATRIX_DETAIL b
    ON a.id = b.header_id
WHERE a.project_id IN
(
    SELECT proj.id
    FROM cmktsch.WF_PROJ proj
    WHERE proj.last_update_date > sysdate - 365
)
AND project_reference_number = :refnum
"""


headers = [
    "Client",
    "Group",
    "Seg",
    "Seg_Offer",
    "Target_Logic",
    "Card_Type",
    "Channel",
    "Lang",
    "T/C",
    "Split",
    "Comm_C",
    "Creative_C",
    "CIS"
]

with get_db_connection() as conn_new:
    with conn_new.cursor() as cursor_new:
        cursor_new.execute(
            sql_new,
            {"refnum": workfront_num}
        )

        row_new = cursor_new.fetchall()


if not row_new:
    df_new = pd.DataFrame(
        columns=headers
    ).set_index("Client")
else:
    df_new = pd.DataFrame(
        row_new,
        columns=headers
    )

    df_new.set_index(
        "Client",
        inplace=True
    )

    df_new_1 = df_new.sort_values(
        by=["Group", "Seg", "Lang"]
    )

    df_new_1["Split"] = (
        df_new_1["Split"].astype(int)
    )


# if not row_new or row_new[0] is None:
#     raise ValueError(
#         f"No data found for Workfront #: {workfront_num}"
#     )

# print(row)
# df_new = pd.DataFrame(row_new, columns=headers)
# df_new.set_index("Client", inplace=True)
# df_new_1 = df_new.sort_values(by=["Group", "Seg", "Lang"])
# df_new_1["Split"] = df_new_1["Split"].astype(int)

segmentation_details = df_new_1


# Excel with suppressions
# (adjust the path correctly)

SP1 = pd.read_excel(
    "Past_Mistakes_20250820.xlsx",
    engine="openpyxl",
    sheet_name="Suppressions"
)

marketing_servicing = df2.loc[
    "Marketing or Servicing",
    "Values"
]

channel = df2.loc[
    "Channel",
    "Values"
]


if marketing_servicing == "Servicing" and channel == "EM":
    standard_suppression = SP1.loc[0, "Query"]

elif marketing_servicing == "Marketing" and channel == "EM":
    standard_suppression = SP1.loc[1, "Query"]

elif marketing_servicing == "Servicing" and channel == "DM":
    standard_suppression = SP1.loc[2, "Query"]

elif marketing_servicing == "Marketing" and channel == "DM":
    standard_suppression = SP1.loc[3, "Query"]

elif marketing_servicing == "Marketing" and channel == "MMSG":
    standard_suppression = SP1.loc[4, "Query"]

else:
    standard_suppression = "Suppressions will be added soon"


# SP2 = pd.read_excel(
#     "Past_Mistakes_20250820.xlsx",
#     engine="openpyxl",
#     sheet_name="Portfolio_Suppressions"
# )

# if marketing_servicing == "Servicing" and channel == "EM":
#     standard_suppression = SP1.loc[0, "Query"]

# elif marketing_servicing == "Marketing" and channel == "EM":
#     standard_suppression = SP1.loc[1, "Query"]

# elif marketing_servicing == "Servicing" and channel == "DM":
#     standard_suppression = SP1.loc[2, "Query"]

# elif marketing_servicing == "Marketing" and channel == "DM":
#     standard_suppression = SP1.loc[3, "Query"]

# elif marketing_servicing == "Marketing" and channel == "MMSG":
#     standard_suppression = SP1.loc[4, "Query"]

DS1 = pd.read_excel(
    "Past_Mistakes_20250820.xlsx",
    engine="openpyxl",
    sheet_name="Past_Mistakes"
)

past_mistakes = ""

for i in range(len(DS1)):
    if DS1.loc[i, "Client"] == df2.loc["Portfolio: Name", "Values"]:
        past_mistakes = DS1.loc[i, "Past_Mistakes"]
        break


# Create Word document

doc = docx.Document()

section = doc.sections[0]

section = doc.sections[0]

paragraph = header_paragraphs[0]
paragraph.text = "TimesNew Roman"
paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


doc.add_heading(
    "Data Sufficiency Report (DSR)",
    0
)

doc.add_heading(
    f"1. Data Availability Percentage: "
    f"{df.loc['Availability_Percentage', 'Values']}%",
    level=1
)

doc.add_heading(
    "2. Analysis Report from Workfront",
    level=1
)

doc.add_paragraph(
    "Missing Data Fields:",
    bold=True
)

t1 = doc.add_table(
    rows=1,
    cols=missing_values.shape[1]
)

t1.style = "Table Grid"

for j in range(missing_values.shape[1]):
    cell = t1.cell(0, j)
    cell.text = missing_values.columns[j]
    set_cell_background(cell, "F4CCC")


for i in range(missing_values.shape[0]):

    row = t1.add_row()

    for j in range(missing_values.shape[1]):

        cell = missing_values.iat[i, j]

        if j == 0:
            set_no_wrap(row.cells[j])

        row.cells[j].text = str(cell)


doc.add_paragraph(" ")

doc.add_paragraph(
    "Populated Data Fields:",
    bold=True
)

t2 = doc.add_table(
    rows=1,
    cols=non_missing_values.shape[1]
)

t2.style = "Table Grid"

for j in range(non_missing_values.shape[1]):

    cell = t2.cell(0, j)
    cell.text = non_missing_values.columns[j]


for i in range(non_missing_values.shape[0]):

    row = t2.add_row()

    for j in range(non_missing_values.shape[1]):

        cell = non_missing_values.iat[i, j]

        if j == 0:
            set_no_wrap(row.cells[j])

        row.cells[j].text = str(cell)


q = doc.add_heading(
    "3. Suppressions List",
    level=1
)

q = doc.add_heading(
    "Standard Suppressions",
    level=2
)

doc.add_paragraph(
    f"{standard_suppression}"
)

q = doc.add_heading(
    "Portfolio Suppressions",
    level=2
)

doc.add_paragraph(
    f"{standard_suppression}"
)

q = doc.add_heading(
    "4. Important Pointers",
    level=1
)

q = doc.add_paragraph(
    f"{past_mistakes}"
)


doc.add_heading(
    "5. Segmentation Matrix",
    level=1
)

p = doc.add_paragraph(
    "Link to Segmentation Matrix: "
)

add_hyperlink(
    p,
    "Link",
    df.loc["Segmentation_Matrix_Link", "Values"]
)


if segmentation_details.empty is False:

    doc.add_paragraph(
        "Segmentation Matrix Preview:",
        bold=True
    )

    t3 = doc.add_table(
        rows=1,
        cols=segmentation_details.shape[1]
    )

    t3.style = "Table Grid"

    for j in range(segmentation_details.shape[1]):

        cell = t3.cell(0, j)
        cell.text = segmentation_details.columns[j]


    for i in range(segmentation_details.shape[0]):

        row = t3.add_row()

        for j in range(segmentation_details.shape[1]):

            cell = segmentation_details.iat[i, j]

            if j == 0:
                set_no_wrap(row.cells[j])

            row.cells[j].text = str(cell)

else:

    doc.add_paragraph(
        "Segmentation Matrix isn't uploaded yet"
    )


filename = "table.docx"

doc.save(filename)

return filename


# ------------------------------------------------------------
# Flask routes and app code
# ------------------------------------------------------------

FORM_HTML = """
<!doctype html>
<title>Workfront Data Sufficiency Report</title>

<style>
body {
    background-color: #FFCE1B;
    display: flex;
    flex-direction: column;
    justify-content: center;
    align-items: center;
    height: 100vh; /* full viewport height */
    margin: 0;
    text-align: center;
    font-family: sans-serif;
}

form {
    margin-top: 1em;
}

ul {
    color: red;
    list-style: none;
    padding: 0;
    margin-top: 1em;
}
</style>

<h1>Data Sufficiency Report (DSR)</h1>

<h2>Please Enter One or More Workfront Number</h2>

<form method="post">
    <input type="text" name="workfront_num" required>
    <input type="submit" value="Generate Report(s)">
</form>

<p>
    Note: Use comma, semicolon, or space to separate multiple numbers.
</p>

<div id="flash-container">
    {% with messages = get_flashed_messages() %}
        {% if messages %}
            <ul style="color: red;">
                {% for message in messages %}
                    <li>{{ message }}</li>
                {% endfor %}
            </ul>
        {% endif %}
    {% endwith %}
</div>

<script>
const form = document.querySelector('form');
const flashContainer = document.getElementById('flash-container');

form.addEventListener('submit', function() {
    // Clear ONLY when the user submits (button click or Enter)
    if (flashContainer) {
        flashContainer.innerHTML = "";
    }
});
</script>
"""


@app.before_request
def limit_remote_addr():
    local_hosts = ['localhost', '127.0.0.1']

    port = ':' + str(
        app.config.get('project_port', '5555')
    )

    local_hosts_ports = [
        host + port for host in local_hosts
    ]

    project_hosts = app.config.get(
        'project_hosts',
        []
    )

    if not project_hosts:
        project_hosts = local_hosts_ports

    if request.host not in project_hosts:
        print(
            "Not allowed host: {}"
            .format(request.host)
        )
        print(
            "Allowed hosts: {}"
            .format(project_hosts)
        )
        abort(403)  # Forbidden


@bp.route("/", methods=["GET", "POST"])
def index():

    if request.method == "POST":

        workfront_num = request.form.get(
            'workfront_num',
            ''
        ).strip()

        if not workfront_num:
            flash(
                "Please enter at least one Workfront number."
            )
            return render_template_string(FORM_HTML)

        # Split input into numbers by common delimiters,
        # filter out bad entries
        raw_entries = [
            x.strip()
            for x in re.split(
                r'[,\s;]+',
                workfront_num
            )
            if x.strip()
        ]

        # Only accept valid 7- or 8-digit entries
        valid_entries = [
            x for x in raw_entries
            if re.fullmatch(r'\d{7,8}', x)
        ]

        # Deduplicate while preserving input order
        seen = set()
        entries = []

        for x in valid_entries:
            if x not in seen:
                seen.add(x)
                entries.append(x)

        if len(entries) < len(valid_entries):
            flash(
                f"Duplicate Workfront numbers were removed. "
                f"Processing {len(entries)} unique number(s)."
            )

        if not entries:
            flash(
                "Please enter at least one valid "
                "7 or 8 digit Workfront number."
            )
            return render_template_string(FORM_HTML)

        # Hold output file paths
        output_files = []

        # Track errors
        errors = []

        for num in entries:

            try:
                filename = generate_report(num)

                # Rename to avoid overwrites
                new_filename = f"DSR_{num}.docx"

                os.rename(
                    filename,
                    new_filename
                )

                output_files.append(
                    new_filename
                )

            except Exception as e:
                errors.append(
                    f"{num}: {e}"
                )

        if errors:
            for msg in errors:
                flash(
                    f"Error for {msg}"
                )

        # If only one report, send as docx
        if len(output_files) == 1:

            return send_file(
                output_files[0],
                as_attachment=True,
                download_name=output_files[0],
                mimetype=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                )
            )

        elif len(output_files) > 1:

            # Multiple: bundle into ZIP
            temp_zip = tempfile.NamedTemporaryFile(
                delete=False,
                suffix=".zip"
            )

            with zipfile.ZipFile(
                temp_zip,
                'w'
            ) as myzip:

                for path in output_files:
                    myzip.write(
                        path,
                        arcname=os.path.basename(path)
                    )

            temp_zip.close()

            return send_file(
                temp_zip.name,
                as_attachment=True,
                download_name="DSR_Reports.zip",
                mimetype="application/zip"
            )

        return render_template_string(
            FORM_HTML
        )


# ------------------------------------------------------------
# Arg parser for the standard Anaconda-project options
# ------------------------------------------------------------

parser = argparse.ArgumentParser(
    prog="imagenet-flask",
    description="Classification Webapp with TensorFlow"
)

parser.add_argument(
    '--anaconda-project-host',
    action='append',
    default=[],
    help='Hostname to allow in requests'
)

parser.add_argument(
    '--anaconda-project-port',
    action='store',
    default=8086,
    type=int,
    help='Port to listen on'
)

parser.add_argument(
    '--anaconda-project-iframe-hosts',
    action='store',
    default=[],
    help=(
        'Space-separated hosts which can embed us in an iframe '
        'per our Content-Security-Policy'
    )
)

parser.add_argument(
    '--anaconda-project-no-browser',
    action='store_true',
    default=False,
    help='Disable opening in a browser'
)

parser.add_argument(
    '--anaconda-project-use-xheaders',
    action='store_true',
    default=False,
    help='Trust X-headers from reverse proxy'
)

parser.add_argument(
    '--anaconda-project-url-prefix',
    action='store',
    default="",
    help='Prefix in front of urls'
)

parser.add_argument(
    '--anaconda-project-address',
    action='store',
    default='0.0.0.0',
    help='IP address the application should listen on.'
)


args = parser.parse_args(sys.argv[1:])

project_hosts = args.anaconda_project_host

app.config['project_hosts'] = project_hosts

app.config['project_port'] = (
    args.anaconda_project_port
)

app.register_blueprint(
    bp,
    url_prefix=args.anaconda_project_url_prefix
)

app.run(
    debug=True,
    port=args.anaconda_project_port,
    host=args.anaconda_project_address
)
