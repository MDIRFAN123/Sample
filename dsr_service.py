"""Backend-only adaptation of the supplied legacy DSR business logic."""
from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any
import tempfile

import docx
import numpy as np
import pandas as pd
import pytz
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

class DSRConfigurationError(RuntimeError):
    pass

@dataclass(frozen=True)
class DSRResult:
    path: Path
    availability_percentage: int

    @property
    def public_data(self) -> dict[str, int]:
        return {"availability_percentage": self.availability_percentage}

RC = {'Data Targeting Criteria':.17,'Marketing or Servicing':.15,'Channel':.14,'Will a start file be provided?':.12,'Is a control group needed?':.11,'Campaign Type - Lifecycle':.09,'Output File Data or MISC fields':.08,'Number of Segments':.06,'Non-Mon Required?':.04,'Marketing Tollgate Approval':.03,'Creative Agency Name':.01}
PS = {**RC,'Marketing Tollgate Approval':.04,'Creative Agency Name':.03,'Non-Mon Required?':.01}
ZRC = {'Data Targeting Criteria':.14,'Marketing or Servicing':.12,'ZETA File Name':.20,'Channel':.11,'Will a start file be provided?':.10,'Is a control group needed?':.09,'Campaign Type - Lifecycle':.07,'Output File Data or MISC fields':.06,'Number of Segments':.05,'Non-Mon Required?':.03,'Marketing Tollgate Approval':.02,'Creative Agency Name':.01}
ZPS = {**ZRC,'Marketing Tollgate Approval':.03,'Creative Agency Name':.02,'Non-Mon Required?':.01}
RC_CLIENTS = {"Vemo Visa","PayPal Co-Brand/PayPal PLCC","Lowe's Commercial","Sam's Club","PayPal Credit","TJX","Cathay Pacific","Verizon","Crate & Barrel All Brands","Fleet Farm (SR6)","ShopH","PayPal All Products","Lifestyle Markets - Multiclients","Acquisitions - B2C","Lowe's Consumer","Dual Card - B2C","Dick's Sporting Goods","At Home","Google","JC Penney","American Signature Furniture","RC Multi-Client","American Eagle","Belk","Harbor Freight Tools","Amazon Consumer","Walgreens","Brand - B2C","Fareportal","PayPal Credit Card","Ebay","J.Crew","Virgin","HSN","Walmart","QVC"}
FIELDS = ("Name","IsStatusComplete","DE:Zeta COPS File Name","DE:Channel","DE:Marketing or Servicing","DE:Campaign Type - Lifecycle","DE:Creative Agency Name","DE:Marketing Tollgate Approval","DE:Will a start file be provided?","DE:Data Targeting Criteria","DE:Output File Data or MISC fields","DE:Is a control group needed?","DE:Number of Segments","DE:COPS Campaign Code","DE:Client or Industry calculated","DE:Non-Mon Required?","DE:Open Segmentation Matrix from Project")
NAMES = ("Name","Is Status Complete?","ZETA File Name","Channel","Marketing or Servicing","Campaign Type - Lifecycle","Creative Agency Name","Marketing Tollgate Approval","Will a start file be provided?","Data Targeting Criteria","Output File Data or MISC fields","Is a control group needed?","Number of Segments","COPS Campaign Code","Portfolio: Name","Non-Mon Required?","Segmentation Matrix_Link")
SEGMENT_COLUMNS = ("Client","Group","Seg","Seg_Offer","Target_Logic","Card_Type","Channel","Lang","T/C","Split","Comm_C","Creative_C","CIS")

class DSRService:
    def __init__(self, database: Any, controlled_workbook: Path) -> None:
        self.database, self.controlled_workbook = database, Path(controlled_workbook)

    def generate_dsr(self, project_reference_number: str, output: Path | None = None, *, task_number: str = "") -> DSRResult:
        ref = str(project_reference_number).strip()
        if not ref: raise ValueError("Project reference number is required.")
        if not self.controlled_workbook.is_file():
            raise DSRConfigurationError("The controlled DSR dependency is unavailable. Contact the administrator.")
        body = self.database.fetch_dsr_project_body(ref)
        params = body.get("parameterValues", {}) if isinstance(body.get("parameterValues"), dict) else {}
        combined = {**{k:v for k,v in body.items() if k in FIELDS}, **{k:v for k,v in params.items() if k in FIELDS}}
        frame = pd.DataFrame({k:combined.get(k,np.nan) for k in FIELDS}, index=["Values"]).T
        frame.index, frame.index.name = NAMES, "Parameters"
        client = "RC" if frame.loc["Portfolio: Name","Values"] in RC_CLIENTS else "PS"
        frame.loc["Client","Values"] = client
        zeta = "zeta" in str(frame.loc["Name","Values"]).casefold()
        weights = (ZRC if client == "RC" else ZPS) if zeta else (RC if client == "RC" else PS)
        availability = int(sum(v for k,v in weights.items() if k in frame.index and not pd.isna(frame.loc[k,"Values"]))*100)
        frame.loc["Availability Percentage","Values"] = availability
        analysis = frame.iloc[:-3] if zeta else frame.drop(index="ZETA File Name").iloc[:-3]
        segments = pd.DataFrame(self.database.fetch_dsr_segmentation(ref), columns=SEGMENT_COLUMNS)
        if not segments.empty:
            segments = segments.set_index("Client").sort_values(["Group","Seg","Lang"])
            segments["Split"] = segments["Split"].astype(int)
        suppression_data = pd.read_excel(self.controlled_workbook, engine="openpyxl", sheet_name="Suppressions")
        row = {("Servicing","EM"):0,("Marketing","EM"):1,("Servicing","DM"):2,("Marketing","DM"):3,("Marketing","MMSG"):4}.get((analysis.loc["Marketing or Servicing","Values"],analysis.loc["Channel","Values"]))
        suppression = suppression_data.loc[row,"Query"] if row is not None else "Suppressions will be added soon"
        mistakes_data = pd.read_excel(self.controlled_workbook, engine="openpyxl", sheet_name="Past_Mistakes")
        matches = mistakes_data.loc[mistakes_data["Client"] == analysis.loc["Portfolio: Name","Values"],"Past_Mistakes"]
        mistakes = "" if matches.empty else matches.iloc[0]
        if output is None:
            temporary = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            temporary.close()
            output = Path(temporary.name)
        else:
            output = Path(output)
        output.parent.mkdir(parents=True, exist_ok=True)
        self._document(output, frame, analysis, segments, suppression, mistakes, ref, str(task_number).strip())
        return DSRResult(output, availability)

    generate = generate_dsr

    @staticmethod
    def _table(document, data, color=None):
        data = data.reset_index(); table = document.add_table(rows=1, cols=len(data.columns)); table.style="Table Grid"
        for j,name in enumerate(data.columns):
            table.cell(0,j).text=str(name)
            if color:
                shd=OxmlElement("w:shd"); shd.set(qn("w:fill"),color); table.cell(0,j)._tc.get_or_add_tcPr().append(shd)
        for values in data.itertuples(index=False,name=None):
            cells=table.add_row().cells
            for j,value in enumerate(values): cells[j].text=str(value)

    @staticmethod
    def _hyperlink(paragraph, url):
        rel=paragraph.part.relate_to(url,docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,is_external=True)
        link=OxmlElement("w:hyperlink"); link.set(qn("r:id"),rel); run=OxmlElement("w:r"); props=OxmlElement("w:rPr")
        color=OxmlElement("w:color"); color.set(qn("w:val"),"0563C1"); underline=OxmlElement("w:u"); underline.set(qn("w:val"),"single")
        props.extend((color,underline)); run.append(props); text=OxmlElement("w:t"); text.text="Link"; run.append(text); link.append(run); paragraph._p.append(link)

    @classmethod
    def _document(cls, output, frame, analysis, segments, suppression, mistakes, project_reference_number, task_number):
        document=docx.Document(); header=document.sections[0].header.paragraphs[0]
        header.text=datetime.now(pytz.timezone("US/Eastern")).strftime("%Y-%m-%d %H:%M:%S"); header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        document.add_heading("Data Sufficiency Report (DSR)",0)
        document.add_paragraph(f"Project Reference Number: {project_reference_number}")
        document.add_paragraph(f"Task Number: {task_number or 'Not provided'}")
        document.add_heading(f"1. Data Availability Percentage: {frame.loc['Availability Percentage','Values']}%",1)
        document.add_heading("2. Analysis Report from Workfront",1)
        document.add_paragraph().add_run("Missing Data Fields:").bold=True; cls._table(document,analysis[analysis["Values"].isnull()].fillna("Missing"),"F4CCCC")
        document.add_paragraph(" "); document.add_paragraph().add_run("Populated Data Fields:").bold=True; cls._table(document,analysis[~analysis["Values"].isnull()].fillna("Missing"))
        document.add_heading("3. Suppressions List",1); document.add_heading("Standard Suppressions",2); document.add_paragraph(str(suppression))
        document.add_heading("Portfolio Suppressions",2); document.add_paragraph(str(suppression))
        document.add_heading("4. Important Pointers",1); document.add_paragraph(str(mistakes))
        document.add_heading("5. Segmentation Matrix",1); paragraph=document.add_paragraph("Link to Segmentation Matrix: "); cls._hyperlink(paragraph,str(frame.loc["Segmentation Matrix_Link","Values"]))
        if segments.empty: document.add_paragraph("Segmentation Matrix isn't uploaded yet")
        else: document.add_paragraph().add_run("Segmentation Matrix Preview:").bold=True; cls._table(document,segments)
        document.save(output)

def generate_dsr(project_reference_number: str, *, database: Any, controlled_workbook: Path, output: Path | None = None, task_number: str = "") -> DSRResult:
    """Generate a DSR without creating or invoking any UI."""
    return DSRService(database, controlled_workbook).generate_dsr(project_reference_number, output, task_number=task_number)
